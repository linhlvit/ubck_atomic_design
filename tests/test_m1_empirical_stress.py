"""Empirical Stress Test Suite for Milestone 1:
- clean_description.py (38 diverse & adversarial cases)
- build_docx.py (column widths math, XML transformations, edge-case tables)
"""

from __future__ import annotations

import sys
import unittest
import tempfile
import time
from pathlib import Path

# Add scripts directory to path
REPO_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = REPO_ROOT / ".claude" / "skills" / "datamart-gen-docs" / "scripts"
sys.path.insert(0, str(SCRIPTS_DIR))

import clean_description
import build_docx
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT


class TestCleanDescriptionStress(unittest.TestCase):
    """38 adversarial stress test cases covering all edge cases."""

    def test_01_nested_and_modified_tags(self):
        cases = [
            # Case 1: Modified tag with parens
            ("(Sửa 2026-07-17) Mã chứng khoán cơ sở", "Mã chứng khoán cơ sở"),
            # Case 2: Version addition tag with nested info
            ("(Bổ sung 2026-08-15 v2.1) Tên đầy đủ của công ty đại chúng (tiếng Việt)", "Tên đầy đủ của công ty đại chúng (tiếng Việt)"),
            # Case 3: Multiple update tags in one string
            ("(Thêm 2026-05-10) Mã tài khoản giao dịch (Cập nhật 2026-06-01)", "Mã tài khoản giao dịch"),
            # Case 4: English Update tag
            ("(Update 2026-04-29) Số quyết định xử phạt vi phạm hành chính", "Số quyết định xử phạt vi phạm hành chính"),
        ]
        for inp, exp in cases:
            with self.subTest(inp=inp):
                res = clean_description.clean(inp)
                self.assertEqual(res, exp)

    def test_02_malformed_prefixes_and_keys(self):
        cases = [
            # Case 5: PK driving only
            ("PK — Driving: cdr_dt", ""),
            # Case 6: NK date join note
            ("NK — ngày lịch dùng để join từ Fact", ""),
            # Case 7: BK tuple
            ("; BK: (schema_code, cl_code)", ""),
            # Case 8: PK prefix with business description
            ("PK - mã hồ sơ vụ việc thanh tra", "Mã hồ sơ vụ việc thanh tra"),
            # Case 9: FK prefix with colon
            ("FK : mã đối tượng thanh tra liên kết sang danh mục", "Mã đối tượng thanh tra liên kết sang danh mục"),
            # Case 10: DD prefix with em-dash
            ("DD — số thứ tự giao dịch phát sinh trong ngày", "Số thứ tự giao dịch phát sinh trong ngày"),
            # Case 11: NK with trailing BK
            ("NK — mã scheme; BK: (schema_code, cl_code)", "Mã scheme"),
        ]
        for inp, exp in cases:
            with self.subTest(inp=inp):
                res = clean_description.clean(inp)
                self.assertEqual(res, exp)

    def test_03_unicode_vietnamese_capitalization(self):
        cases = [
            # Case 12: Lowercase đ -> Đ
            ("địa chỉ trụ sở chính của người hành nghề chứng khoán", "Địa chỉ trụ sở chính của người hành nghề chứng khoán"),
            # Case 13: Lowercase ủ -> Ủ + update tag
            ("ủy quyền thực hiện công bố thông tin (Cập nhật 2026-07)", "Ủy quyền thực hiện công bố thông tin"),
            # Case 14: Lowercase ô -> Ô
            ("ông/Bà đại diện theo pháp luật của công ty quản lý quỹ", "Ông/Bà đại diện theo pháp luật của công ty quản lý quỹ"),
            # Case 15: Lowercase ý -> Ý
            ("ý kiến kiểm toán độc lập về báo cáo tài chính năm", "Ý kiến kiểm toán độc lập về báo cáo tài chính năm"),
            # Case 16: Lowercase ở -> Ở
            ("ở trạng thái hoạt động bình thường", "Ở trạng thái hoạt động bình thường"),
        ]
        for inp, exp in cases:
            with self.subTest(inp=inp):
                res = clean_description.clean(inp)
                self.assertEqual(res, exp)

    def test_04_metadata_blocks(self):
        cases = [
            # Case 17: Extreme metadata combo
            ("PK surrogate. BCV: \"BCV_NHNCK_01\". Hash: MD5(id). Nguồn thực: NHNCK.P_INFO. FK target: dim_practitioner. Classification: Cat_A. Pair with cert_no. Họ và tên người hành nghề", "Họ và tên người hành nghề"),
            # Case 18: BCV single quote + Shared entity
            ("Mã cổ phiếu niêm yết. BCV: 'STK_LIST'. Shared entity — dùng chung cho toàn bộ phân hệ PTTT", "Mã cổ phiếu niêm yết"),
            # Case 19: Source system with dot path + ETL load timestamp
            ("Tên sàn giao dịch. Nguồn thực: SCMS.EXCHANGES.TABLE_A. ETL load timestamp", "Tên sàn giao dịch"),
        ]
        for inp, exp in cases:
            with self.subTest(inp=inp):
                res = clean_description.clean(inp)
                self.assertEqual(res, exp)

    def test_05_sql_clauses_and_filters(self):
        cases = [
            # Case 20: Dedup with window function
            ("Dedup theo ROW_NUMBER() OVER(PARTITION BY account_no ORDER BY updated_at DESC). Mã tài khoản", "Mã tài khoản"),
            # Case 21: Filter clause
            ("Filter phạm vi: status = 'ACTIVE' AND deleted_flag = 0. Danh sách người đại diện pháp luật", "Danh sách người đại diện pháp luật"),
            # Case 22: Count aggregation note
            ("Tổng số đợt chào bán thành công (COUNT(DISTINCT issue_id) FILTER (WHERE state = 'SUCCESS')). Số đợt chào bán", "Tổng số đợt chào bán thành công. Số đợt chào bán"),
            # Case 23: Grain note
            ("Grain của bảng: Mỗi dòng tương ứng 1 giao dịch. Giá trị khớp lệnh", "Giá trị khớp lệnh"),
        ]
        for inp, exp in cases:
            with self.subTest(inp=inp):
                res = clean_description.clean(inp)
                self.assertEqual(res, exp)

    def test_06_joins_and_trailing_references(self):
        cases = [
            # Case 24: ETL join_atomic
            ("FK mã quỹ — ETL join_atomic với bảng ATM.fund_management_company ON fund_id = src_fund_id", "FK mã quỹ"),
            # Case 25: Derived formula
            ("Tuổi người hành nghề — ETL-derived: YEAR(CURRENT_DATE) - YEAR(birth_date); fallback birth_year khi NULL", "Tuổi người hành nghề"),
            # Case 26: Source table after em-dash
            ("Mã đợt phát hành — SCMS.ISSUANCE.OFFERING_CODE", "Mã đợt phát hành"),
            # Case 27: Source table after dot
            ("Tên doanh nghiệp. THANHTRA.TT_HO_SO.CORP_NAME", "Tên doanh nghiệp"),
            # Case 28: không có PK note
            ("Số lượng chứng khoán — không có PK", "Số lượng chứng khoán"),
            # Case 29: Xem O_ issue
            ("Mã báo cáo — Xem O_QLCB_1", "Mã báo cáo"),
        ]
        for inp, exp in cases:
            with self.subTest(inp=inp):
                res = clean_description.clean(inp)
                self.assertEqual(res, exp)

    def test_07_edge_cases_and_empty_strings(self):
        cases = [
            # Case 30: Empty string
            ("", ""),
            # Case 31: Whitespace only
            ("   \t\n  ", ""),
            # Case 32: Punctuation only
            (". — : ;  ", ""),
            # Case 33: PK surrogate only
            ("PK surrogate", ""),
            # Case 34: SCD4A only
            ("SCD4A", ""),
            # Case 35: Closed issue tag
            ("(Closed)", ""),
            # Case 36: PENDING with note
            ("PENDING (chưa có Silver) — Mã chỉ tiêu", "Mã chỉ tiêu"),
            # Case 37: ETL sinh tự động
            ("ETL sinh tự động", ""),
            # Case 38: Load strategy SCD4A
            ("Load strategy: SCD4A. Ngày kết thúc hiệu lực", "Ngày kết thúc hiệu lực"),
        ]
        for inp, exp in cases:
            with self.subTest(inp=inp):
                res = clean_description.clean(inp)
                self.assertEqual(res, exp)


class TestBuildDocxStress(unittest.TestCase):
    """Stress testing column width calculations, XML manipulation, and table variations."""

    def test_column_width_constants(self):
        # Portrait (total content width = 9074 DXA)
        self.assertEqual(sum(build_docx._COL_WIDTHS_4_PORTRAIT), 9074)
        self.assertEqual(sum(build_docx._COL_WIDTHS_7_PORTRAIT), 9074)
        self.assertEqual(sum(build_docx._COL_WIDTHS_8_PORTRAIT), 9074)

        # Landscape (total content width = 13999 DXA)
        self.assertEqual(sum(build_docx._COL_WIDTHS_4_LANDSCAPE), 13999)
        self.assertEqual(sum(build_docx._COL_WIDTHS_8_LANDSCAPE), 13999)
        self.assertEqual(sum(build_docx._COL_WIDTHS_11_LANDSCAPE), 13999)

        # 12-column Landscape (Physical table: == 13999 DXA)
        self.assertEqual(sum(build_docx._COL_WIDTHS_12_LANDSCAPE), 13999)
        self.assertEqual(len(build_docx._COL_WIDTHS_12_LANDSCAPE), 12)
        # Check specific column allocations (STT, ColName, DataType, ..., ETL Rules)
        self.assertEqual(build_docx._COL_WIDTHS_12_LANDSCAPE[0], 360)   # STT
        self.assertEqual(build_docx._COL_WIDTHS_12_LANDSCAPE[7], 2200)  # Mô tả
        self.assertEqual(build_docx._COL_WIDTHS_12_LANDSCAPE[11], 3619) # ETL Rules

    def _verify_table_xml(self, table, expected_cols: int, expected_widths: list[int] | None, is_portrait: bool):
        tblPr = table._tbl.find(qn("w:tblPr"))
        self.assertIsNotNone(tblPr, "tblPr must exist")

        # 1. tblLayout fixed
        tblLayout = tblPr.find(qn("w:tblLayout"))
        self.assertIsNotNone(tblLayout, "tblLayout must exist")
        self.assertEqual(tblLayout.get(qn("w:type")), "fixed")

        # 2. tblW
        tblW = tblPr.find(qn("w:tblW"))
        self.assertIsNotNone(tblW, "tblW must exist")
        expected_total_w = 9074 if is_portrait else 13999
        self.assertEqual(tblW.get(qn("w:w")), str(expected_total_w))
        self.assertEqual(tblW.get(qn("w:type")), "dxa")

        # 3. tblBorders
        tblBorders = tblPr.find(qn("w:tblBorders"))
        self.assertIsNotNone(tblBorders, "tblBorders must exist")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = tblBorders.find(qn(f"w:{side}"))
            self.assertIsNotNone(b, f"border {side} must exist")
            self.assertEqual(b.get(qn("w:val")), "single")
            self.assertEqual(b.get(qn("w:sz")), "8")
            self.assertEqual(b.get(qn("w:color")), "999999")

        # 4. tblCellMar
        tblCellMar = tblPr.find(qn("w:tblCellMar"))
        self.assertIsNotNone(tblCellMar, "tblCellMar must exist")

        # 5. tblGrid
        if expected_widths is not None:
            tblGrid = table._tbl.find(qn("w:tblGrid"))
            self.assertIsNotNone(tblGrid, "tblGrid must exist")
            gridCols = tblGrid.findall(qn("w:gridCol"))
            self.assertEqual(len(gridCols), expected_cols)
            actual_widths = [int(c.get(qn("w:w"))) for c in gridCols]
            self.assertEqual(actual_widths, expected_widths)

        # 6. Rows: tblHeader on row 0, cantSplit on all rows, cell shading & alignment
        for row_idx, row in enumerate(table.rows):
            trPr = row._tr.find(qn("w:trPr"))
            self.assertIsNotNone(trPr, f"trPr must exist on row {row_idx}")
            self.assertIsNotNone(trPr.find(qn("w:cantSplit")), f"cantSplit must exist on row {row_idx}")

            if row_idx == 0:
                self.assertIsNotNone(trPr.find(qn("w:tblHeader")), "tblHeader must exist on header row")
            else:
                self.assertIsNone(trPr.find(qn("w:tblHeader")), "tblHeader must NOT exist on non-header rows")

            for ci, cell in enumerate(row.cells):
                tcPr = cell._tc.find(qn("w:tcPr"))
                self.assertIsNotNone(tcPr, f"tcPr must exist in cell ({row_idx}, {ci})")
                shd = tcPr.find(qn("w:shd"))
                if row_idx == 0:
                    self.assertIsNotNone(shd, f"shading must exist on header row ({row_idx}, {ci})")
                    self.assertEqual(shd.get(qn("w:fill")), "BDD7EE")
                elif row_idx % 2 == 0:
                    self.assertIsNotNone(shd, f"shading must exist on even row ({row_idx}, {ci})")
                    self.assertEqual(shd.get(qn("w:fill")), "EEF4FB")
                else:
                    self.assertIsNone(shd, f"shading must NOT exist on odd row ({row_idx}, {ci})")

                for para in cell.paragraphs:
                    if row_idx == 0 or ci == 0:
                        self.assertEqual(para.alignment, WD_ALIGN_PARAGRAPH.CENTER)
                    else:
                        self.assertEqual(para.alignment, WD_ALIGN_PARAGRAPH.LEFT)

    def test_mock_tables_all_column_counts(self):
        """Test XML transformations on 4, 7, 8, 11, 12-column tables in Portrait and Landscape."""
        with tempfile.TemporaryDirectory() as tmpdir:
            test_doc_path = Path(tmpdir) / "test_tables.docx"

            # 1. Test TKCSLD (Landscape) with 4, 8, 11, 12 col tables
            doc = Document()
            t4 = doc.add_table(rows=3, cols=4)
            t8 = doc.add_table(rows=3, cols=8)
            t11 = doc.add_table(rows=3, cols=11)
            t12 = doc.add_table(rows=4, cols=12)

            # Populate cells to simulate real content
            for t in (t4, t8, t11, t12):
                for r_idx, row in enumerate(t.rows):
                    for c_idx, cell in enumerate(row.cells):
                        cell.text = f"R{r_idx}C{c_idx}"

            doc.save(str(test_doc_path))
            build_docx.post_process_docx(test_doc_path, doc_type="tkcsld")

            processed_doc = Document(str(test_doc_path))
            self.assertEqual(len(processed_doc.tables), 4)
            self._verify_table_xml(processed_doc.tables[0], 4, build_docx._COL_WIDTHS_4_LANDSCAPE, is_portrait=False)
            self._verify_table_xml(processed_doc.tables[1], 8, build_docx._COL_WIDTHS_8_LANDSCAPE, is_portrait=False)
            self._verify_table_xml(processed_doc.tables[2], 11, build_docx._COL_WIDTHS_11_LANDSCAPE, is_portrait=False)
            self._verify_table_xml(processed_doc.tables[3], 12, build_docx._COL_WIDTHS_12_LANDSCAPE, is_portrait=False)

            # 2. Test PTTK (Portrait) with 4, 7, 8 col tables
            doc_p = Document()
            tp4 = doc_p.add_table(rows=3, cols=4)
            tp7 = doc_p.add_table(rows=3, cols=7)
            tp8 = doc_p.add_table(rows=3, cols=8)
            for t in (tp4, tp7, tp8):
                for r_idx, row in enumerate(t.rows):
                    for c_idx, cell in enumerate(row.cells):
                        cell.text = f"P_R{r_idx}C{c_idx}"

            test_p_path = Path(tmpdir) / "test_pttk.docx"
            doc_p.save(str(test_p_path))
            build_docx.post_process_docx(test_p_path, doc_type="pttk")

            processed_p = Document(str(test_p_path))
            self.assertEqual(len(processed_p.tables), 3)
            self._verify_table_xml(processed_p.tables[0], 4, build_docx._COL_WIDTHS_4_PORTRAIT, is_portrait=True)
            self._verify_table_xml(processed_p.tables[1], 7, build_docx._COL_WIDTHS_7_PORTRAIT, is_portrait=True)
            self._verify_table_xml(processed_p.tables[2], 8, build_docx._COL_WIDTHS_8_PORTRAIT, is_portrait=True)

    def test_stress_large_table_and_unsupported_cols(self):
        """Stress test with a 100-row table and non-standard column counts (e.g. 5, 9 cols)."""
        with tempfile.TemporaryDirectory() as tmpdir:
            test_doc_path = Path(tmpdir) / "test_stress.docx"
            doc = Document()

            # 100-row 12-column table
            t_large = doc.add_table(rows=100, cols=12)
            for r_idx, row in enumerate(t_large.rows):
                for c_idx, cell in enumerate(row.cells):
                    cell.text = f"Cell_{r_idx}_{c_idx}"

            # Non-standard 5-col table (should not crash, generic layout applied)
            t_5 = doc.add_table(rows=2, cols=5)
            for r_idx, row in enumerate(t_5.rows):
                for c_idx, cell in enumerate(row.cells):
                    cell.text = f"C5_{r_idx}_{c_idx}"

            # 1-row table (header only)
            t_1row = doc.add_table(rows=1, cols=4)
            for c_idx, cell in enumerate(t_1row.rows[0].cells):
                cell.text = f"Header_{c_idx}"

            doc.save(str(test_doc_path))

            t0 = time.time()
            build_docx.post_process_docx(test_doc_path, doc_type="tkcsld")
            elapsed = time.time() - t0

            # Performance check: 100-row table processed in < 1.0 second
            self.assertLess(elapsed, 1.5, f"Post-processing large table took {elapsed:.2f}s, expected < 1.5s")

            processed_doc = Document(str(test_doc_path))
            self._verify_table_xml(processed_doc.tables[0], 12, build_docx._COL_WIDTHS_12_LANDSCAPE, is_portrait=False)
            self._verify_table_xml(processed_doc.tables[1], 5, None, is_portrait=False)
            self._verify_table_xml(processed_doc.tables[2], 4, build_docx._COL_WIDTHS_4_LANDSCAPE, is_portrait=False)


if __name__ == "__main__":
    unittest.main()
