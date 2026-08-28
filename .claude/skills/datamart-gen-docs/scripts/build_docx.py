"""Phase 3: Convert Markdown đã approved → DOCX cho tài liệu PTTK hoặc TKCSLD.

Spec tài liệu:
- Landscape A4 (16834 x 11909 DXA)
- Font: Times New Roman
- Bảng 4/8/11 cột theo độ rộng chuẩn (tổng = 13999 DXA)
- Header row: shading #BDD7EE, đậm, căn giữa
- Hàng chẵn (data): shading #EEF4FB
- Diagram mermaid → PNG (dùng mmdc)

Usage:
    # DOCX từng module riêng lẻ
    python build_docx.py --module TT --type pttk
    python build_docx.py --module TT --type tkcsld
    python build_docx.py --module TT --type both

    # DOCX tổng (sau khi chạy merge_md.py)
    python build_docx.py --module ALL --type pttk
    python build_docx.py --module ALL --type tkcsld

Yêu cầu:
    pip install python-docx
    pandoc >= 3.0  (https://pandoc.org/installing.html)
    npm install -g @mermaid-js/mermaid-cli   # để render mermaid → PNG
"""

from __future__ import annotations

import argparse
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path

# Ensure Pandoc in PATH if present in local appdata or common install locations
pandoc_appdata = Path(os.environ.get("LOCALAPPDATA", "C:/Users/ADMIN/AppData/Local")) / "Pandoc"
if pandoc_appdata.exists() and str(pandoc_appdata) not in os.environ.get("PATH", ""):
    os.environ["PATH"] = str(pandoc_appdata) + os.pathsep + os.environ.get("PATH", "")

# Ensure npm global binaries (mmdc) in PATH
npm_appdata = Path(os.environ.get("APPDATA", "C:/Users/ADMIN/AppData/Roaming")) / "npm"
if npm_appdata.exists() and str(npm_appdata) not in os.environ.get("PATH", ""):
    os.environ["PATH"] = str(npm_appdata) + os.pathsep + os.environ.get("PATH", "")

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")

REPO_ROOT = Path(__file__).resolve().parents[4]
OUTPUT_DIR = REPO_ROOT / "docs" / "output" / "datamart"


def _module_dir(module: str) -> Path:
    # File tổng (ALL hoặc tên custom từ merge_md.py) nằm thẳng ở OUTPUT_DIR
    if module.upper() == "ALL" or "_" in module:
        return OUTPUT_DIR
    return OUTPUT_DIR / module


def _md_path(module: str, doc_type: str) -> Path:
    return _module_dir(module) / f"DTM_{module}_{doc_type.upper()}.md"


def _docx_path(module: str, doc_type: str) -> Path:
    return _module_dir(module) / f"DTM_{module}_{doc_type.upper()}.docx"


# ─── Mermaid → PNG ───────────────────────────────────────────────────────────

def render_mermaid_blocks(md_text: str, out_dir: Path) -> tuple[str, list[Path]]:
    """Tìm tất cả ```mermaid...``` trong MD, render thành PNG, thay bằng ![](path)."""
    mmdc = shutil.which("mmdc")
    if not mmdc:
        print(
            "[WARN] mmdc không tìm thấy — diagram giữ nguyên code block trong DOCX.\n"
            "       Cài đặt: npm install -g @mermaid-js/mermaid-cli",
            file=sys.stderr,
        )
        return md_text, []

    png_files: list[Path] = []
    pattern = re.compile(r"```mermaid\n(.*?)```", re.DOTALL)
    counter = [0]

    def replace(m: re.Match) -> str:
        counter[0] += 1
        mmd_code = m.group(1)
        mmd_file = out_dir / f"_diagram_{counter[0]}.mmd"
        png_file  = out_dir / f"_diagram_{counter[0]}.png"
        mmd_file.write_text(mmd_code, encoding="utf-8")

        result = subprocess.run(
            [mmdc, "-i", str(mmd_file), "-o", str(png_file), "-b", "white"],
            capture_output=True, text=True,
        )
        mmd_file.unlink(missing_ok=True)
        if result.returncode != 0:
            print(f"[WARN] mmdc lỗi diagram {counter[0]}: {result.stderr[:200]}", file=sys.stderr)
            return m.group(0)

        png_files.append(png_file)
        return f"![Diagram {counter[0]}]({png_file.name})"

    new_text = pattern.sub(replace, md_text)
    return new_text, png_files


# ─── DOCX post-process ───────────────────────────────────────────────────────

# Độ rộng cột (DXA) — landscape A4, content width = 13999
_COL_WIDTHS_4_LANDSCAPE  = [400, 3500, 2700, 7399]
_COL_WIDTHS_8_LANDSCAPE  = [360, 2000, 1500, 560, 560, 560, 680, 7779]
_COL_WIDTHS_11_LANDSCAPE = [360, 1500, 1200, 480, 480, 480, 680, 2519, 1700, 1700, 2900]
_COL_WIDTHS_12_LANDSCAPE = [360, 1300, 1100, 480, 480, 480, 480, 2200, 900, 1400, 1200, 3619]

# Độ rộng cột (DXA) — portrait A4, content width = 9074
_COL_WIDTHS_4_PORTRAIT  = [360, 2500, 2000, 4214]
_COL_WIDTHS_7_PORTRAIT  = [360, 1500, 1200, 600, 1000, 1000, 3414]
_COL_WIDTHS_8_PORTRAIT  = [360, 1500, 1200, 450, 450, 450, 550, 4114]

_HEADER_SHADING   = "BDD7EE"
_EVEN_ROW_SHADING = "EEF4FB"


def _apply_col_widths(table, widths: list[int], qn_fn, OxmlElement) -> None:
    if len(table.columns) != len(widths):
        return
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            tcPr = cell._tc.find(qn_fn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                cell._tc.insert(0, tcPr)
            tcW = tcPr.find(qn_fn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn_fn("w:w"), str(widths[ci]))
            tcW.set(qn_fn("w:type"), "dxa")


def _set_shading(cell, color: str, qn_fn, OxmlElement) -> None:
    tcPr = cell._tc.find(qn_fn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        cell._tc.insert(0, tcPr)
    shd = tcPr.find(qn_fn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn_fn("w:val"), "clear")
    shd.set(qn_fn("w:color"), "auto")
    shd.set(qn_fn("w:fill"), color)


def _find_reference_doc(doc_type: str, custom_ref: Path | None = None) -> Path | None:
    """Tìm kiếm file template mẫu --reference-doc theo thứ tự ưu tiên."""
    if custom_ref and custom_ref.exists():
        return custom_ref

    candidates = [
        # 1. Bản mẫu Q5 hiện hành (đặt tại gốc thư mục skill) — ưu tiên cao nhất cho PTTK
        REPO_ROOT / ".claude" / "skills" / "datamart-gen-docs" / "UBCKNN_Q5_Tai lieu phan tich thiet ke_v1.0_20260429.docx",
        # 2. Thư mục reference trong skill
        REPO_ROOT / ".claude" / "skills" / "datamart-gen-docs" / "reference" / "UBCKNN_Q5_Tai lieu phan tich thiet ke_v1.0_20260429.docx",
        REPO_ROOT / ".claude" / "skills" / "datamart-gen-docs" / "reference" / "UBCKNN_Q5_Tai lieu phan tich thiet ke_Template.docx",
        REPO_ROOT / ".claude" / "skills" / "datamart-gen-docs" / "reference" / "UBCKNN_Thiet ke co so du lieu_Template.docx",
        # 3. Thư mục template sample của dự án (fallback)
        REPO_ROOT / "docs" / "templates" / "sample" / "UBCKNN_Q5_Tai lieu phan tich thiet ke_Template.docx",
    ]
    if doc_type.lower() == "tkcsld":
        candidates.insert(
            0,
            REPO_ROOT / ".claude" / "skills" / "datamart-gen-docs" / "reference" / "UBCKNN_Thiet ke co so du lieu_Template.docx"
        )
        candidates.insert(
            1,
            REPO_ROOT / "docs" / "templates" / "sample" / "UBCKNN_Thiet ke co so du lieu_Template.docx"
        )

    for p in candidates:
        if p.exists():
            return p
    return None


def post_process_docx(docx_path: Path, doc_type: str = "pttk") -> None:
    """Post-process: Portrait (PTTK) / Landscape (TKCSLD), Times New Roman, bảng theo spec."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.shared import Pt, Emu

    doc = Document(str(docx_path))
    is_portrait = (doc_type.lower() == "pttk")

    # ── Page Setup ──
    # Portrait: W=11909, H=16834 DXA; content_w = 9074 DXA
    # Landscape: W=16834, H=11909 DXA; content_w = 13999 DXA
    page_w_dxa = 11909 if is_portrait else 16834
    page_h_dxa = 16834 if is_portrait else 11909
    table_w_dxa = 9074 if is_portrait else 13999

    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT if is_portrait else WD_ORIENT.LANDSCAPE
        pgSz = section._sectPr.find(qn("w:pgSz"))
        if pgSz is None:
            pgSz = OxmlElement("w:pgSz")
            section._sectPr.append(pgSz)
        pgSz.set(qn("w:w"), str(page_w_dxa))
        pgSz.set(qn("w:h"), str(page_h_dxa))
        pgSz.set(qn("w:orient"), "portrait" if is_portrait else "landscape")
        
        # 1 DXA = 914400/1440 EMU = 635 EMU
        section.page_width      = Emu(page_w_dxa * 635)
        section.page_height     = Emu(page_h_dxa * 635)
        section.left_margin     = Emu(1701 * 635)
        section.right_margin    = Emu(1134 * 635)
        section.top_margin      = Emu(1418 * 635)
        section.bottom_margin   = Emu(1418 * 635)
        section.header_distance = Emu(720 * 635)
        section.footer_distance = Emu(720 * 635)

    # ── Font body paragraphs ──
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = "Times New Roman"
            if run.font.size is None:
                run.font.size = Pt(12)

    # ── Fix bảng ──
    for table in doc.tables:
        table.autofit = False
        tblPr = table._tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            table._tbl.insert(0, tblPr)

        # 1. Fixed layout
        for old in tblPr.findall(qn("w:tblLayout")):
            tblPr.remove(old)
        tblLayout = OxmlElement("w:tblLayout")
        tblLayout.set(qn("w:type"), "fixed")
        tblPr.append(tblLayout)

        # 2. Table width = full page content
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:w"), str(table_w_dxa))
        tblW.set(qn("w:type"), "dxa")

        # 3. Borders — SINGLE, size=1 (8 eighth-pt = 1pt), color=#999999
        for old in tblPr.findall(qn("w:tblBorders")):
            tblPr.remove(old)
        tblBorders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "8")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "999999")
            tblBorders.append(el)
        tblPr.append(tblBorders)

        # 4. Cell margin — top/bottom=60, left/right=100 DXA
        for old in tblPr.findall(qn("w:tblCellMar")):
            tblPr.remove(old)
        tblCellMar = OxmlElement("w:tblCellMar")
        for side, val in [("top", "60"), ("left", "100"), ("bottom", "60"), ("right", "100")]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"), val)
            el.set(qn("w:type"), "dxa")
            tblCellMar.append(el)
        tblPr.append(tblCellMar)

        # 5. Apply column widths theo số cột và nhận dạng loại bảng
        ncols = len(table.columns)
        col_widths = None

        if ncols == 4:
            col_widths = _COL_WIDTHS_4_PORTRAIT if is_portrait else _COL_WIDTHS_4_LANDSCAPE
        elif ncols == 7:
            col_widths = _COL_WIDTHS_7_PORTRAIT
        elif ncols == 8:
            col_widths = _COL_WIDTHS_8_PORTRAIT if is_portrait else _COL_WIDTHS_8_LANDSCAPE
        elif ncols == 11:
            col_widths = _COL_WIDTHS_11_LANDSCAPE
        elif ncols == 12:
            col_widths = _COL_WIDTHS_12_LANDSCAPE

        if col_widths and len(col_widths) == ncols:
            _apply_col_widths(table, col_widths, qn, OxmlElement)

            # Update tblGrid
            old_grid = table._tbl.find(qn("w:tblGrid"))
            if old_grid is not None:
                table._tbl.remove(old_grid)
            tblGrid = OxmlElement("w:tblGrid")
            for w in col_widths:
                gridCol = OxmlElement("w:gridCol")
                gridCol.set(qn("w:w"), str(w))
                tblGrid.append(gridCol)
            tblPr_idx = table._tbl.index(tblPr)
            table._tbl.insert(tblPr_idx + 1, tblGrid)

        # 6. Shading + Header repeat + cantSplit
        for row_idx, row in enumerate(table.rows):
            is_header = (row_idx == 0)
            is_even   = not is_header and (row_idx % 2 == 0)

            trPr = row._tr.get_or_add_trPr()
            # CantSplit
            if trPr.find(qn("w:cantSplit")) is None:
                trPr.append(OxmlElement("w:cantSplit"))
            if is_header:
                if trPr.find(qn("w:tblHeader")) is None:
                    trPr.append(OxmlElement("w:tblHeader"))

            for ci, cell in enumerate(row.cells):
                if is_header:
                    _set_shading(cell, _HEADER_SHADING, qn, OxmlElement)
                elif is_even:
                    _set_shading(cell, _EVEN_ROW_SHADING, qn, OxmlElement)

                for para in cell.paragraphs:
                    # Alignment: header row → center; cột STT (ci=0) → center
                    if is_header or ci == 0:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # Xóa indent thừa
                    pPr = para._p.get_or_add_pPr()
                    for ind in pPr.findall(qn("w:ind")):
                        pPr.remove(ind)
                    ind_el = OxmlElement("w:ind")
                    ind_el.set(qn("w:left"), "0")
                    ind_el.set(qn("w:firstLine"), "0")
                    pPr.append(ind_el)
                    # spacing before/after = 0
                    spacing = pPr.find(qn("w:spacing"))
                    if spacing is None:
                        spacing = OxmlElement("w:spacing")
                        pPr.append(spacing)
                    spacing.set(qn("w:before"), "0")
                    spacing.set(qn("w:after"), "0")

                    # Vertical align center
                    tcPr = cell._tc.find(qn("w:tcPr"))
                    if tcPr is None:
                        tcPr = OxmlElement("w:tcPr")
                        cell._tc.insert(0, tcPr)
                    vAlign = tcPr.find(qn("w:vAlign"))
                    if vAlign is None:
                        vAlign = OxmlElement("w:vAlign")
                        tcPr.append(vAlign)
                    vAlign.set(qn("w:val"), "center")

                    for run in para.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)  # 20 half-pt
                        if is_header:
                            run.font.bold = True

    doc.save(str(docx_path))
    orient_str = "portrait A4" if is_portrait else "landscape A4"
    print(
        f"Post-process done — {orient_str}, Times New Roman, shading, col widths: {docx_path}",
        file=sys.stderr,
    )


def _sanitize_docx_package(docx_path: Path) -> None:
    """Ensure [Content_Types].xml has required extensions (e.g. ttf, odttf) to prevent KeyError in python-docx."""
    with zipfile.ZipFile(docx_path, "r") as zin:
        entries = {item.filename: zin.read(item.filename) for item in zin.infolist()}

    if "[Content_Types].xml" in entries:
        ct = entries["[Content_Types].xml"].decode("utf-8")
        dirty = False
        if 'Extension="ttf"' not in ct and "Extension='ttf'" not in ct:
            ct = ct.replace("</Types>", '<Default Extension="ttf" ContentType="application/x-font-ttf"/></Types>')
            dirty = True
        if dirty:
            entries["[Content_Types].xml"] = ct.encode("utf-8")
            with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
                for name, data in entries.items():
                    zout.writestr(name, data)


# ─── Main build ──────────────────────────────────────────────────────────────

def build(module: str, doc_type: str, reference_doc: Path | None = None) -> Path:
    md_file = _md_path(module, doc_type)
    if not md_file.exists():
        raise SystemExit(
            f"ERROR: Không tìm thấy {md_file}\n"
            "Hãy chạy Phase 1 trước (Claude gen Markdown) và đảm bảo đã được phê duyệt."
        )

    out_dir = _module_dir(module)
    out_dir.mkdir(parents=True, exist_ok=True)

    # Bước 1: render mermaid → PNG
    md_text = md_file.read_text(encoding="utf-8")
    md_text_replaced, png_files = render_mermaid_blocks(md_text, out_dir)

    # Bước 2: ghi temp MD
    tmp_md = Path(tempfile.mktemp(suffix=".md"))
    tmp_md.write_text(md_text_replaced, encoding="utf-8")

    # Bước 3: kiểm tra pandoc
    pandoc = shutil.which("pandoc")
    if not pandoc:
        tmp_md.unlink(missing_ok=True)
        raise SystemExit(
            "ERROR: pandoc không tìm thấy.\n"
            "macOS:   brew install pandoc\n"
            "Windows: winget install JohnMacFarlane.Pandoc"
        )

    # Bước 4: pandoc MD → DOCX
    tmp_docx = Path(tempfile.mktemp(suffix=".docx"))
    ref_docx = _find_reference_doc(doc_type, reference_doc)
    cmd = [
        pandoc,
        str(tmp_md),
        "--from", "markdown+pipe_tables+auto_identifiers",
        "--to", "docx",
        "--output", str(tmp_docx),
        "--resource-path", str(out_dir),
    ]
    if ref_docx:
        cmd.extend(["--reference-doc", str(ref_docx)])

    try:
        result = subprocess.run(cmd, capture_output=True, text=True)
    finally:
        tmp_md.unlink(missing_ok=True)

    if result.returncode != 0:
        tmp_docx.unlink(missing_ok=True)
        raise SystemExit(f"ERROR Pandoc:\n{result.stderr}")
    if result.stderr:
        print(result.stderr, file=sys.stderr)

    # Bước 4.5: sanitize [Content_Types].xml nếu có embedded fonts từ template
    _sanitize_docx_package(tmp_docx)

    # Bước 5: post-process
    post_process_docx(tmp_docx, doc_type=doc_type)

    # Bước 6: copy sang đích
    out_docx = _docx_path(module, doc_type)
    shutil.copy2(tmp_docx, out_docx)
    tmp_docx.unlink(missing_ok=True)

    # Dọn temp PNG
    for png in png_files:
        png.unlink(missing_ok=True)

    print(f"OK: {out_docx}")
    return out_docx


def main() -> None:
    parser = argparse.ArgumentParser(
        description=__doc__,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("--module", required=True, help="Mã module: TT, NHNCK, NDTNN, QLCB, ...")
    parser.add_argument(
        "--type", required=True, choices=["pttk", "tkcsld", "both"],
        help="Loại tài liệu: pttk | tkcsld | both",
    )
    parser.add_argument(
        "--reference-doc", type=Path, default=None,
        help="Đường dẫn tùy chỉnh tới file template DOCX reference",
    )
    args = parser.parse_args()

    types = ["pttk", "tkcsld"] if args.type == "both" else [args.type]
    for t in types:
        build(args.module.upper(), t, reference_doc=args.reference_doc)


if __name__ == "__main__":
    main()
