"""Script cập nhật trực tiếp 10 phân hệ Datamart vào tài liệu chuẩn UBCKNN Q5
(UBCKNN_Q5_Tai lieu phan tich thiet ke_v1.0_20260429.docx) bảo toàn 100% template,
bìa, chữ ký, kiến trúc hệ thống, API, Web App mockup và styling chuẩn mực.
"""

from __future__ import annotations

import io
import os
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")

REPO_ROOT = Path(__file__).resolve().parents[4]
SKILL_DIR = REPO_ROOT / ".claude" / "skills" / "datamart-gen-docs"
OUTPUT_DIR = REPO_ROOT / "docs" / "output" / "datamart"
MASTER_DOCX_PATH = OUTPUT_DIR / "UBCKNN_Q5_Tai_lieu_phan_tich_thiet_ke_v1.0_20260429.docx"

# Ensure mmdc in PATH
npm_appdata = Path(os.environ.get("APPDATA", "C:/Users/ADMIN/AppData/Roaming")) / "npm"
if npm_appdata.exists() and str(npm_appdata) not in os.environ.get("PATH", ""):
    os.environ["PATH"] = str(npm_appdata) + os.pathsep + os.environ.get("PATH", "")

MODULES = [
    ("TT", "1", "HOẠT ĐỘNG THANH TRA"),
    ("NHNCK", "2", "NGƯỜI HÀNH NGHỀ CHỨNG KHOÁN"),
    ("NDTNN", "3", "NHÀ ĐẦU TƯ NƯỚC NGOÀI"),
    ("QLCB", "4", "QUẢN LÝ CHÀO BÁN"),
    ("GSDC", "5", "GIÁM SÁT CÔNG TY ĐẠI CHÚNG"),
    ("GSTT", "6", "GIÁM SÁT THỊ TRƯỜNG"),
    ("QLQ", "7", "CÔNG TY QUẢN LÝ QUỸ (AMC)"),
    ("QLKD", "8", "HOẠT ĐỘNG CÔNG TY CHỨNG KHOÁN"),
    ("PTTT", "9", "PHÂN TÍCH THỊ TRƯỜNG"),
    ("TKNB", "10", "THỐNG KÊ THỊ TRƯỜNG"),
]


def render_mermaid(mmd_code: str, out_png: Path) -> bool:
    """Render 1 Mermaid code block sang PNG dùng mmdc với độ nét cao."""
    out_png = out_png.resolve()
    out_png.parent.mkdir(parents=True, exist_ok=True)
    if out_png.exists() and out_png.stat().st_size > 0:
        return True

    mmdc = shutil.which("mmdc")
    if not mmdc:
        if (npm_appdata / "mmdc.cmd").exists():
            mmdc = str(npm_appdata / "mmdc.cmd")
        elif (npm_appdata / "mmdc.CMD").exists():
            mmdc = str(npm_appdata / "mmdc.CMD")
        elif (npm_appdata / "mmdc").exists():
            mmdc = str(npm_appdata / "mmdc")

    if not mmdc:
        print("[WARN] mmdc not found, diagram will be skipped", file=sys.stderr)
        return False

    with tempfile.NamedTemporaryFile(suffix=".mmd", delete=False, mode="w", encoding="utf-8") as f:
        f.write(mmd_code)
        tmp_mmd = Path(f.name)

    try:
        res = subprocess.run(
            [str(mmdc), "-i", str(tmp_mmd), "-o", str(out_png), "-b", "white", "-s", "2"],
            shell=True,
            capture_output=True, text=True
        )
        if res.returncode != 0:
            print(f"[WARN] mmdc render error for {out_png.name}: {res.stderr}", file=sys.stderr)
        return res.returncode == 0 and out_png.exists() and out_png.stat().st_size > 0
    finally:
        tmp_mmd.unlink(missing_ok=True)


def parse_bullet_item(raw_text: str) -> tuple[str, str]:
    """Phân tách chính xác tên Entity (in đậm) và phần giải thích (thường) từ 1 dòng bullet."""
    raw_text = raw_text.strip()
    
    # Case 1: **Entity Name:** Description OR **Entity Name**: Description
    m = re.match(r"^\*{2}(.*?)\*{2}:?\s*(.*)$", raw_text)
    if m:
        ent = m.group(1).strip().rstrip(":")
        desc = m.group(2).strip().lstrip(":")
        desc = re.sub(r"^\*{1,2}\s*", "", desc)
        return ent, desc
    
    # Case 2: Entity Name: Description
    m2 = re.match(r"^(.*?):\s*(.*)$", raw_text)
    if m2:
        ent = m2.group(1).replace("*", "").strip()
        desc = m2.group(2).replace("**", "").strip()
        return ent, desc
    
    return "", raw_text.replace("**", "").strip()


def parse_module_pttk(md_path: Path, temp_img_dir: Path, mod_code: str, mod_idx: str, mod_name: str) -> list[dict]:
    """Phân tích cú pháp 1 file Markdown PTTK thành danh sách các phần tử dữ liệu chuẩn hóa."""
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    elements = []
    
    # 1. Heading cho Module: 3.1.X LUỒNG ĐỒNG BỘ DỮ LIỆU CHO NHÓM BÁO CÁO ...
    elements.append({
        "type": "heading3",
        "text": f"3.1.{mod_idx}  LUỒNG ĐỒNG BỘ DỮ LIỆU CHO NHÓM BÁO CÁO {mod_name}"
    })

    # 2. Heading: 3.1.X.1 Thông tin chung luồng đồng bộ
    elements.append({
        "type": "heading4",
        "text": f"3.1.{mod_idx}.1 Thông tin chung luồng đồng bộ"
    })

    in_info = False
    in_workflow = False
    subgroups = []
    current_subgroup = None

    idx = 0
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        if re.search(r"^###\s*3\.[12]\.\d+\.1\b", stripped) or (stripped.startswith("###") and "Thông tin chung" in stripped):
            in_info = True
            in_workflow = False
            idx += 1
            continue

        if re.search(r"^###\s*3\.[12]\.\d+\.2\b", stripped) or (stripped.startswith("###") and "Luồng nghiệp vụ" in stripped):
            in_info = False
            in_workflow = True
            elements.append({
                "type": "heading4",
                "text": f"3.1.{mod_idx}.2 Luồng nghiệp vụ"
            })
            idx += 1
            continue

        if in_info:
            if stripped.startswith("- ") or stripped.startswith("* "):
                raw_bullet = re.sub(r"^[-*]\s+", "", stripped).strip()
                ent_name, desc_body = parse_bullet_item(raw_bullet)
                elements.append({
                    "type": "info_item",
                    "label": f"{ent_name}: " if ent_name else "",
                    "value": desc_body
                })
            idx += 1
            continue

        if in_workflow:
            if stripped.startswith("#### "):
                group_title = stripped.replace("####", "").strip()
                group_title = re.sub(r"^3\.[12]\.\d+\.2\.\d+\s*", "", group_title)
                current_subgroup_idx = len(subgroups) + 1
                full_group_heading = f"3.1.{mod_idx}.2.{current_subgroup_idx} {group_title}"
                
                current_subgroup = {
                    "heading": full_group_heading,
                    "diagram_code": None,
                    "purpose": None,
                    "staging_atomic": [],
                    "atomic_datamart": []
                }
                subgroups.append(current_subgroup)
                idx += 1
                continue

            if stripped.startswith("```mermaid"):
                mmd_lines = []
                idx += 1
                while idx < len(lines) and not lines[idx].strip().startswith("```"):
                    mmd_lines.append(lines[idx])
                    idx += 1
                mmd_code = "\n".join(mmd_lines).strip()
                if current_subgroup:
                    current_subgroup["diagram_code"] = mmd_code
                idx += 1
                continue

            if "**Mục đích:**" in stripped or "Mục đích:" in stripped:
                purpose_text = re.sub(r"^\*{0,2}Mục đích:\*{0,2}\s*", "", stripped).strip()
                purpose_text = purpose_text.replace("**", "").strip()
                if current_subgroup:
                    current_subgroup["purpose"] = purpose_text
                idx += 1
                continue

            if "Staging → Atomic:" in stripped or "Staging -> Atomic:" in stripped:
                idx += 1
                while idx < len(lines):
                    l_sub = lines[idx].strip()
                    if "Atomic → Datamart:" in l_sub or "Atomic -> Datamart:" in l_sub or l_sub.startswith("####") or l_sub.startswith("###") or l_sub.startswith("##"):
                        break
                    if l_sub.startswith("- ") or l_sub.startswith("* "):
                        clean_item = re.sub(r"^[-*]\s+", "", l_sub).strip()
                        if current_subgroup:
                            current_subgroup["staging_atomic"].append(clean_item)
                    idx += 1
                continue

            if "Atomic → Datamart:" in stripped or "Atomic -> Datamart:" in stripped:
                idx += 1
                while idx < len(lines):
                    l_sub = lines[idx].strip()
                    if l_sub.startswith("####") or l_sub.startswith("###") or l_sub.startswith("##"):
                        break
                    if l_sub.startswith("- ") or l_sub.startswith("* "):
                        clean_item = re.sub(r"^[-*]\s+", "", l_sub).strip()
                        if current_subgroup:
                            current_subgroup["atomic_datamart"].append(clean_item)
                    idx += 1
                continue

        idx += 1

    # Thêm subgroups vào elements
    for sg_idx, sg in enumerate(subgroups, 1):
        elements.append({
            "type": "heading5",
            "text": sg["heading"]
        })
        
        # Render diagram nếu có
        if sg["diagram_code"]:
            img_path = temp_img_dir / f"diagram_{mod_code}_{sg_idx}.png"
            if not (img_path.exists() and img_path.stat().st_size > 0):
                render_mermaid(sg["diagram_code"], img_path)
            if img_path.exists() and img_path.stat().st_size > 0:
                elements.append({
                    "type": "image",
                    "path": str(img_path.resolve())
                })

        # Mục đích
        if sg["purpose"]:
            elements.append({
                "type": "purpose",
                "label": "Mục đích: ",
                "text": sg["purpose"]
            })

        # Mô tả luồng
        elements.append({
            "type": "flow_desc_header",
            "text": "Mô tả luồng:"
        })

        if sg["staging_atomic"]:
            elements.append({
                "type": "sub_header",
                "text": "Staging → Atomic:"
            })
            for item in sg["staging_atomic"]:
                elements.append({
                    "type": "desc_item",
                    "text": item
                })

        if sg["atomic_datamart"]:
            elements.append({
                "type": "sub_header",
                "text": "Atomic → Datamart:"
            })
            for item in sg["atomic_datamart"]:
                elements.append({
                    "type": "desc_item",
                    "text": item
                })

    return elements


def format_paragraph(p, style_name="Normal", space_before=Pt(0), space_after=Pt(3), line_spacing=1.15, left_indent=Pt(0), first_line_indent=Pt(0), align=None, keep_with_next=False):
    """Thiết lập chính xác các thông số lề và khoảng cách cho đoạn văn bản, triệt tiêu hoàn toàn hanging indent từ style gốc."""
    p.style = style_name
    pf = p.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing = line_spacing
    pf.left_indent = left_indent
    pf.first_line_indent = first_line_indent
    pf.keep_with_next = keep_with_next
    if align is not None:
        p.alignment = align
    
    # Ghi đè trực tiếp thẻ w:ind trong XML của đoạn để chặn hoàn toàn thuộc tính hanging âm từ style gốc
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    
    left_dxa = int(left_indent.pt * 20) if left_indent is not None else 0
    first_dxa = int(first_line_indent.pt * 20) if first_line_indent is not None else 0
    
    ind.set(qn("w:left"), str(left_dxa))
    if first_dxa < 0:
        ind.set(qn("w:hanging"), str(-first_dxa))
        ind.attrib.pop(qn("w:firstLine"), None)
    else:
        ind.set(qn("w:firstLine"), str(first_dxa))
        ind.attrib.pop(qn("w:hanging"), None)


def add_text_run(p, text, font_name="Times New Roman", font_size=Pt(12), bold=False, italic=False, color=None):
    """Thêm một Run với font Times New Roman và định dạng đồng nhất."""
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    return run


def update_master_q5_document():
    print(f"=== BẮT ĐẦU CẬP NHẬT TÀI LIỆU CHUẨN UBCKNN Q5 ===")
    print(f"Template gốc: {MASTER_DOCX_PATH} ({MASTER_DOCX_PATH.stat().st_size:,} bytes)")

    if not MASTER_DOCX_PATH.exists():
        raise FileNotFoundError(f"Không tìm thấy template chuẩn tại {MASTER_DOCX_PATH}")

    temp_img_dir = OUTPUT_DIR / "temp_q5_images"
    temp_img_dir.mkdir(parents=True, exist_ok=True)

    # 1. Parse toàn bộ 10 module PTTK
    all_module_elements = []
    total_diagrams = 0
    for mod_code, mod_idx, mod_name in MODULES:
        md_file = OUTPUT_DIR / mod_code / f"DTM_{mod_code}_PTTK.md"
        print(f"Parsing module {mod_code} ({mod_name})...")
        elems = parse_module_pttk(md_file, temp_img_dir, mod_code, mod_idx, mod_name)
        img_count = sum(1 for e in elems if e["type"] == "image")
        total_diagrams += img_count
        print(f"  -> {mod_code}: {len(elems)} elements, {img_count} diagrams rendered.")
        all_module_elements.append((mod_code, elems))

    print(f"\nTổng số lưu đồ đã sinh và sẵn sàng chèn: {total_diagrams}/90 diagrams.")

    # 2. Đọc file DOCX gốc vào memory để tránh bị lock khi đọc
    with open(MASTER_DOCX_PATH, "rb") as f:
        doc_stream = io.BytesIO(f.read())
    doc = docx.Document(doc_stream)

    # 3. Định vị vị trí 3.1 trong tài liệu gốc
    p_etl_idx = None
    p_api_idx = None

    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Phân hệ ETL" and p.style.name.startswith("Heading"):
            p_etl_idx = i
        if p.text.strip().startswith("Phân hệ tích hợp dữ liệu") and p.style.name.startswith("Heading"):
            p_api_idx = i
            break

    if p_etl_idx is None or p_api_idx is None:
        raise ValueError(f"Không xác định được ranh giới Section 3.1: p_etl={p_etl_idx}, p_api={p_api_idx}")

    print(f"Vị trí Section 3.1: từ paragraph {p_etl_idx+1} đến {p_api_idx-1} (tổng {p_api_idx - p_etl_idx - 1} paragraphs cũ)")

    # Chuẩn hóa paragraph tiêu đề 'Phân hệ ETL' để thẳng hàng tuyệt đối với lề trái
    p_etl = doc.paragraphs[p_etl_idx]
    p_etl_pPr = p_etl._p.get_or_add_pPr()
    p_etl_ind = p_etl_pPr.find(qn("w:ind"))
    if p_etl_ind is not None:
        p_etl_pPr.remove(p_etl_ind)

    target_anchor = doc.paragraphs[p_api_idx]

    # 4. Chèn toàn bộ nội dung mới của 10 phân hệ trước target_anchor với styling chuẩn mực
    print("Đang chèn 10 phân hệ Datamart chuẩn hóa vào Section 3.1...")
    
    for mod_code, elems in all_module_elements:
        for el in elems:
            el_type = el["type"]
            
            # --- Heading 3: Module (3.1.x LUỒNG ĐỒNG BỘ DỮ LIỆU CHO NHÓM BÁO CÁO...) ---
            if el_type == "heading3":
                p = target_anchor.insert_paragraph_before("", style="Heading 3")
                format_paragraph(p, style_name="Heading 3", space_before=Pt(14), space_after=Pt(6), line_spacing=1.15, left_indent=Pt(0), first_line_indent=Pt(0), keep_with_next=True)
                add_text_run(p, el["text"], font_size=Pt(13), bold=True, italic=False)
            
            # --- Heading 4: Section (3.1.x.1 Thông tin chung / 3.1.x.2 Luồng nghiệp vụ) ---
            elif el_type == "heading4":
                p = target_anchor.insert_paragraph_before("", style="Heading 4")
                format_paragraph(p, style_name="Heading 4", space_before=Pt(10), space_after=Pt(4), line_spacing=1.15, left_indent=Pt(0), first_line_indent=Pt(0), keep_with_next=True)
                add_text_run(p, el["text"], font_size=Pt(12), bold=True, italic=False)
            
            # --- Info items under 3.1.x.1 ---
            elif el_type == "info_item":
                p = target_anchor.insert_paragraph_before("", style="Normal")
                format_paragraph(p, style_name="Normal", space_before=Pt(1), space_after=Pt(2.5), line_spacing=1.15, left_indent=Pt(18), first_line_indent=Pt(-14))
                add_text_run(p, "•  ", font_size=Pt(12), bold=False)
                if el["label"]:
                    add_text_run(p, el["label"], font_size=Pt(12), bold=True)
                if el["value"]:
                    add_text_run(p, el["value"], font_size=Pt(12), bold=False)

            # --- Heading 5: Subgroup (3.1.x.2.y Nhóm thông tin...) ---
            elif el_type == "heading5":
                p = target_anchor.insert_paragraph_before("", style="Heading 5")
                format_paragraph(p, style_name="Heading 5", space_before=Pt(10), space_after=Pt(4), line_spacing=1.15, left_indent=Pt(0), first_line_indent=Pt(0), keep_with_next=True)
                add_text_run(p, el["text"], font_size=Pt(12), bold=True, italic=True)
            
            # --- Diagram image ---
            elif el_type == "image":
                p = target_anchor.insert_paragraph_before("", style="Normal")
                format_paragraph(p, style_name="Normal", space_before=Pt(6), space_after=Pt(6), left_indent=Pt(0), first_line_indent=Pt(0), align=WD_ALIGN_PARAGRAPH.CENTER)
                run = p.add_run()
                try:
                    run.add_picture(el["path"], width=Inches(6.2))
                except Exception as e:
                    print(f"[WARN] Không thể chèn ảnh {el['path']}: {e}", file=sys.stderr)
            
            # --- Purpose ---
            elif el_type == "purpose":
                p = target_anchor.insert_paragraph_before("", style="Normal")
                format_paragraph(p, style_name="Normal", space_before=Pt(4), space_after=Pt(4), line_spacing=1.15, left_indent=Pt(0), first_line_indent=Pt(0))
                add_text_run(p, el["label"], font_size=Pt(12), bold=True)
                add_text_run(p, el["text"], font_size=Pt(12), bold=False)
            
            # --- Flow description header ---
            elif el_type == "flow_desc_header":
                p = target_anchor.insert_paragraph_before("", style="Normal")
                format_paragraph(p, style_name="Normal", space_before=Pt(6), space_after=Pt(2), line_spacing=1.15, left_indent=Pt(0), first_line_indent=Pt(0), keep_with_next=True)
                add_text_run(p, el["text"], font_size=Pt(12), bold=True)
            
            # --- Sub header (Staging -> Atomic / Atomic -> Datamart) ---
            elif el_type == "sub_header":
                p = target_anchor.insert_paragraph_before("", style="Normal")
                format_paragraph(p, style_name="Normal", space_before=Pt(4), space_after=Pt(2), line_spacing=1.15, left_indent=Pt(0), first_line_indent=Pt(0), keep_with_next=True)
                add_text_run(p, el["text"], font_size=Pt(12), bold=True, italic=True)
            
            # --- Desc item (Table mapping bullet) ---
            elif el_type == "desc_item":
                p = target_anchor.insert_paragraph_before("", style="Normal")
                format_paragraph(p, style_name="Normal", space_before=Pt(1), space_after=Pt(2.5), line_spacing=1.15, left_indent=Pt(18), first_line_indent=Pt(-14))
                add_text_run(p, "•  ", font_size=Pt(12), bold=False)
                ent_name, desc_body = parse_bullet_item(el["text"])
                if ent_name:
                    add_text_run(p, f"{ent_name}: ", font_size=Pt(12), bold=True)
                if desc_body:
                    add_text_run(p, desc_body, font_size=Pt(12), bold=False)

    # 5. Xóa các paragraph cũ từ p_etl_idx + 1 đến p_api_idx - 1
    print("Đang xóa các đoạn văn bản cũ của 3.1 cũ...")
    paragraphs_to_remove = doc.paragraphs[p_etl_idx + 1 : p_api_idx]
    for p in paragraphs_to_remove:
        p._p.getparent().remove(p._p)

    # 6. Lưu file kết quả
    out_docx_1 = OUTPUT_DIR / "UBCKNN_Q5_Tai_lieu_phan_tich_thiet_ke_v1.0_20260429.docx"
    out_docx_fallback = OUTPUT_DIR / "UBCKNN_Q5_Tai_lieu_phan_tich_thiet_ke_v1.1_20260429.docx"
    
    saved_target = out_docx_1
    try:
        doc.save(str(out_docx_1))
        print(f"\nSaved: {out_docx_1} ({out_docx_1.stat().st_size:,} bytes)")
    except PermissionError:
        doc.save(str(out_docx_fallback))
        saved_target = out_docx_fallback
        print(f"\n[INFO] File {out_docx_1.name} đang mở trong Word. Đã lưu sang bản cập nhật: {out_docx_fallback.name} ({out_docx_fallback.stat().st_size:,} bytes)")

    print(f"\n=== HOÀN TẤT CẬP NHẬT TÀI LIỆU CHUẨN Q5 ===")
    print(f"File xuất bản: {saved_target} ({saved_target.stat().st_size:,} bytes)")


if __name__ == "__main__":
    update_master_q5_document()
