"""CLI sinh tài liệu Word bàn giao Atomic Lakehouse — Thiết kế CSDL.

Modes:
    sample  — render 1-2 entity đại diện cho 1 source ra file MD review nhanh
    full    — render fragment đầy đủ cho 1 source vào docs/output/fragments/
    docx    — gom tất cả fragment + render master MD + convert sang DOCX bằng Pandoc

Usage:
    python gen_handover.py --source FIMS --mode sample
    python gen_handover.py --source FIMS --mode full
    python gen_handover.py --mode docx
    python gen_handover.py --mode docx --source FIMS   # chỉ 1 source

Yêu cầu cho mode docx:
    pandoc >= 3.0  (https://pandoc.org/installing.html)
"""

from __future__ import annotations

import argparse
import shutil
import subprocess
import sys
import tempfile
from datetime import datetime
from pathlib import Path

# Thêm scripts/ vào sys.path để import data_loader, render_md, filters
SCRIPTS_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(SCRIPTS_DIR))

import data_loader as DL
import render_md as RM


# Repo root = tổ tiên cách 4 cấp: scripts/ → atomic-gen-docs/ → skills/ → .claude/ → repo
REPO_ROOT = SCRIPTS_DIR.parents[3]
TEMPLATES_DIR = REPO_ROOT / "docs" / "templates"
OUTPUT_DIR = REPO_ROOT / "docs" / "output"
FRAGMENTS_DIR = OUTPUT_DIR / "fragments"  # legacy — không dùng trực tiếp nữa


def _source_dir(source: str) -> Path:
    return OUTPUT_DIR / source


def _fragments_dir(source: str) -> Path:
    return _source_dir(source) / "fragments"


def _sample_dir(source: str) -> Path:
    return _source_dir(source) / "sample"


def cmd_sample(source: str, sample_count: int = 2) -> Path:
    """Render sample MD cho 1 source. Output: docs/output/{SOURCE}/sample/sample_dbdesign.md."""
    source_data = DL.load_source(REPO_ROOT, source, sample=True, sample_count=sample_count)
    if not source_data["entities"]:
        raise SystemExit(f"ERROR: Không tìm thấy entity nào cho source '{source}' trong manifest.csv")

    env = RM.make_env(TEMPLATES_DIR)
    fragment = RM.render_fragment(env, source_data, idx=1)

    out_dir = _sample_dir(source)
    out_dir.mkdir(parents=True, exist_ok=True)

    md_file = out_dir / "sample_dbdesign.md"
    md_file.write_text(fragment, encoding="utf-8")

    uid_groups = source_data.get("uid_groups", [])
    dbml_file = out_dir / f"{source}.dbml"
    dbml_file.write_text(DL.build_dbml(source, source_data["entities"], uid_groups), encoding="utf-8")

    print(f"DBML: {dbml_file}", file=sys.stderr)
    return md_file


def cmd_full(source: str) -> Path:
    """Render fragment đầy đủ cho 1 source. Idempotent — chỉ ghi đè file của source đó.

    Sinh đồng thời vào docs/output/{SOURCE}/fragments/:
    - {SOURCE}_dbdesign.md
    - {SOURCE}.dbml  (tổng hợp)
    - {SOURCE}_UID*.dbml  (per UID group)
    """
    source_data = DL.load_source(REPO_ROOT, source, sample=False)
    if not source_data["entities"]:
        raise SystemExit(f"ERROR: Không tìm thấy entity nào cho source '{source}' trong manifest.csv")

    env = RM.make_env(TEMPLATES_DIR)
    # idx tạm = {IDX}; idx thực sẽ inject lại ở mode docx khi gom master
    fragment = RM.render_fragment(env, source_data, idx="{IDX}")

    frag_dir = _fragments_dir(source)
    frag_dir.mkdir(parents=True, exist_ok=True)

    md_file = frag_dir / f"{source}_dbdesign.md"
    md_file.write_text(fragment, encoding="utf-8")

    uid_groups = source_data.get("uid_groups", [])
    all_entities = source_data["entities"]

    # DBML tổng hợp với TableGroup blocks
    dbml_file = frag_dir / f"{source}.dbml"
    dbml_file.write_text(DL.build_dbml(source, all_entities, uid_groups), encoding="utf-8")
    print(f"DBML (full): {dbml_file}", file=sys.stderr)

    # DBML riêng theo từng UID group
    for g in uid_groups:
        uid_entities = [e for e in all_entities if e["atomic_entity"] in g["entities"]]
        if not uid_entities:
            continue
        safe_project = f"{source}_{g['uid']}"
        uid_dbml = DL.build_dbml(safe_project, uid_entities, note=f"{source} — {g['uid']} {g['label']}")
        uid_file = frag_dir / f"{source}_{g['uid']}.dbml"
        uid_file.write_text(uid_dbml, encoding="utf-8")
        print(f"DBML ({g['uid']}): {uid_file}", file=sys.stderr)

    return md_file


def _load_sources_order() -> list[str] | None:
    """Đọc docs/output/sources_order.txt nếu có. 1 source/dòng, bỏ qua dòng trống/comment."""
    order_file = OUTPUT_DIR / "sources_order.txt"
    if not order_file.exists():
        return None
    sources: list[str] = []
    for line in order_file.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        sources.append(line)
    return sources


def _build_master_md(source_filter: str | None = None) -> tuple[Path, list[str]]:
    """Gom fragment, render master MD. Trả về (md_file, danh sách sources đã gom)."""
    # Tìm tất cả fragment files trong {SOURCE}/fragments/{SOURCE}_dbdesign.md
    if source_filter:
        frag_file = _fragments_dir(source_filter) / f"{source_filter}_dbdesign.md"
        if not frag_file.exists():
            raise SystemExit(f"ERROR: Không tìm thấy fragment cho source '{source_filter}'. Chạy '--mode full --source {source_filter}' trước.")
        discovered = [source_filter]
    else:
        discovered = sorted(
            d.name for d in OUTPUT_DIR.iterdir()
            if d.is_dir() and (d / "fragments" / f"{d.name}_dbdesign.md").exists()
        )
        if not discovered:
            raise SystemExit(f"ERROR: Không có fragment nào trong {OUTPUT_DIR}. Chạy '--mode full --source X' trước.")

    order = _load_sources_order()
    if order:
        sources = [s for s in order if s in discovered]
        sources.extend(s for s in discovered if s not in sources)
    else:
        sources = discovered

    sources_data: list[dict] = []
    fragments_text: list[str] = []
    all_classifications: list[dict] = []
    all_pendings: list[dict] = []

    for idx, source in enumerate(sources, start=1):
        sd = DL.load_source(REPO_ROOT, source, sample=False)
        sources_data.append(sd)
        frag_text = (_fragments_dir(source) / f"{source}_dbdesign.md").read_text(encoding="utf-8")
        frag_text = frag_text.replace("{IDX}", str(idx))
        fragments_text.append(frag_text)
        for c in sd["classifications"]:
            all_classifications.append({**c, "source": source})
        all_pendings.extend(sd["pendings"])

    env = RM.make_env(TEMPLATES_DIR)
    master_md = RM.render_master(
        env=env,
        sources_data=sources_data,
        sources=sources,
        fragments=fragments_text,
        all_classifications=all_classifications,
        all_pendings=all_pendings,
        generated_at=datetime.now().strftime("%Y-%m-%d"),
    )

    # Master MD đặt trong {SOURCE}/ (single source) hoặc OUTPUT_DIR/ (multi)
    if source_filter:
        out_dir = _source_dir(source_filter)
    else:
        out_dir = OUTPUT_DIR
    out_dir.mkdir(parents=True, exist_ok=True)
    suffix = f"_{source_filter}" if source_filter else ""
    master_file = out_dir / f"atomic_dbdesign{suffix}.md"
    master_file.write_text(master_md, encoding="utf-8")
    return master_file, sources


def cmd_docx(source_filter: str | None = None) -> Path:
    """Build master MD rồi convert sang DOCX bằng Pandoc + reference-doc template UBCK."""
    # Bước 1 — build master MD
    master_file, sources = _build_master_md(source_filter)
    print(f"Master MD: {master_file}", file=sys.stderr)
    print(f"Sources: {', '.join(sources)}", file=sys.stderr)

    # Bước 2 — kiểm tra Pandoc (PATH trước, sau đó thử vị trí cài mặc định trên Windows)
    pandoc = shutil.which("pandoc")
    if not pandoc:
        _win_default = Path.home() / "AppData" / "Local" / "Pandoc" / "pandoc.exe"
        if _win_default.exists():
            pandoc = str(_win_default)
        else:
            print(
                "\nERROR: Pandoc không tìm thấy.\n"
                "Cài đặt: winget install JohnMacFarlane.Pandoc\n"
                f"\nMaster MD đã sẵn sàng tại: {master_file}",
                file=sys.stderr,
            )
            raise SystemExit(1)

    # Bước 3 — convert MD → DOCX
    ref_docx = TEMPLATES_DIR / "sample" / "UBCKNN_Thiet ke co so du lieu_Template.docx"
    suffix = f"_{source_filter}" if source_filter else ""
    # Single source → {SOURCE}/; multi source → output root
    docx_dir = _source_dir(source_filter) if source_filter else OUTPUT_DIR
    docx_dir.mkdir(parents=True, exist_ok=True)
    out_docx = docx_dir / f"Atomic_DBDesign{suffix}.docx"

    # Pandoc output ra temp để tránh permission denied khi file đang mở trong Word/OneDrive
    tmp_out = Path(tempfile.mktemp(suffix=".docx"))
    cmd = [
        pandoc,
        str(master_file),
        "--from", "markdown+pipe_tables+auto_identifiers",
        "--to", "docx",
        "--output", str(tmp_out),
    ]
    # Copy reference-doc ra temp để tránh OneDrive/Word lock
    tmp_ref = None
    if ref_docx.exists():
        tmp_ref = Path(tempfile.mktemp(suffix=".docx"))
        shutil.copy2(ref_docx, tmp_ref)
        cmd += ["--reference-doc", str(tmp_ref)]
        print(f"Reference doc: {ref_docx}", file=sys.stderr)
    else:
        print(f"WARNING: Template không tìm thấy tại {ref_docx} — sinh DOCX không có style.", file=sys.stderr)

    try:
        result = subprocess.run(cmd, capture_output=True, text=True)
    finally:
        if tmp_ref and tmp_ref.exists():
            tmp_ref.unlink(missing_ok=True)
    if result.returncode != 0:
        print(f"ERROR Pandoc:\n{result.stderr}", file=sys.stderr)
        tmp_out.unlink(missing_ok=True)
        raise SystemExit(result.returncode)

    if result.stderr:
        print(result.stderr, file=sys.stderr)

    # Bước 4 — set landscape + fix table styling (trên temp file)
    _post_process_docx(tmp_out)

    # Bước 5 — copy temp → đích (ghi đè dù file đang mở read-only qua OneDrive)
    try:
        shutil.copy2(tmp_out, out_docx)
    finally:
        tmp_out.unlink(missing_ok=True)

    return out_docx


def _post_process_docx(docx_path: Path) -> None:
    """Post-process DOCX: landscape A4 + fix table cell styling (font 9pt, no indent, header bold)."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Mm, Pt

    doc = Document(str(docx_path))

    # --- Landscape A4, margin 15mm ---
    for section in doc.sections:
        section.page_width = Mm(297)
        section.page_height = Mm(210)
        pgSz = section._sectPr.find(qn("w:pgSz"))
        if pgSz is None:
            pgSz = OxmlElement("w:pgSz")
            section._sectPr.append(pgSz)
        pgSz.set(qn("w:orient"), "landscape")
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)

    # --- Fix indent + style cho tất cả para ngoài bảng ---
    _HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4",
                       "Heading 5", "Heading 6"}
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        # Heading: giữ nguyên
        if style_name in _HEADING_STYLES:
            continue
        # List Paragraph (bullet): giữ numPr, chỉ fix bold inheritance + keepNext
        if "List" in style_name:
            pPr2 = para._p.get_or_add_pPr()
            kn = pPr2.find(qn("w:keepNext"))
            if kn is None:
                kn = OxmlElement("w:keepNext")
                pPr2.append(kn)
            for run in para.runs:
                if run.font.bold is None:
                    run.font.bold = False
            continue
        # Mọi para khác (Body Text, Block Text, Normal, v.v.): xóa numPr thừa + reset indent + đổi về Normal
        pPr = para._p.get_or_add_pPr()
        for numPr in pPr.findall(qn("w:numPr")):
            pPr.remove(numPr)
        for ind in pPr.findall(qn("w:ind")):
            pPr.remove(ind)
        ind_el = OxmlElement("w:ind")
        ind_el.set(qn("w:left"), "0")
        ind_el.set(qn("w:firstLine"), "0")
        pPr.append(ind_el)
        if style_name not in ("Normal", ""):
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None:
                pStyle.set(qn("w:val"), "Normal")
            else:
                pStyle = OxmlElement("w:pStyle")
                pStyle.set(qn("w:val"), "Normal")
                pPr.insert(0, pStyle)

    # Column widths (twips) cho các loại bảng theo số cột — lấy từ template UBCK
    # 4 cột: STT | Thực thể | Tên bảng | Mô tả
    _COL_WIDTHS_4 = [643, 3132, 2700, 6570]
    # 12 cột: STT | Tên trường | Tên cột | Kiểu DL | Nullable | Unique | P/F Key | Mặc định | Mô tả | Schema.Table | Source Field | ETL Rules
    # Cột "Tên cột" mở rộng (+1440 twips) thay cho cột "Hệ thống nguồn" đã bỏ
    _COL_WIDTHS_12 = [720, 1440, 2440, 1114, 720, 720, 720, 990, 1800, 1170, 1440, 2520]
    # 2 cột PK: Tên trường | Tên cột
    _COL_WIDTHS_PK = [2700, 2700]
    # 5 cột FK: Tên trường | Tên cột | Bảng tham chiếu | Trường tham chiếu | Cột tham chiếu
    _COL_WIDTHS_FK = [1800, 1400, 2200, 2000, 1400]

    def _apply_col_widths(table, widths: list[int]) -> None:
        if len(table.columns) != len(widths):
            return
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                tcPr = cell._tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    cell._tc.insert(0, tcPr)
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is None:
                    tcW = OxmlElement("w:tcW")
                    tcPr.append(tcW)
                tcW.set(qn("w:w"), str(widths[ci]))
                tcW.set(qn("w:type"), "dxa")

    # --- Fix tất cả bảng: cell margin nhỏ, font 9pt, header row bold, column widths ---
    for table in doc.tables:
        tblPr = table._tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            table._tbl.insert(0, tblPr)

        # Set table width = auto (Word tự tính từ tcW của từng cột)
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:w"), "0")
        tblW.set(qn("w:type"), "auto")

        # Override tblCellMar để ghi đè style "Normal Table" (left=108 twips)
        for old in tblPr.findall(qn("w:tblCellMar")):
            tblPr.remove(old)
        tblCellMar = OxmlElement("w:tblCellMar")
        for side in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"), "30")   # 30 twips ≈ 0.5mm
            el.set(qn("w:type"), "dxa")
            tblCellMar.append(el)
        tblPr.append(tblCellMar)

        # Set table borders — single 0.5pt
        for old in tblPr.findall(qn("w:tblBorders")):
            tblPr.remove(old)
        tblBorders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")      # 4 eighth-points = 0.5pt
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
            tblBorders.append(el)
        tblPr.append(tblBorders)

        # Apply column widths theo số cột
        ncols = len(table.columns)
        header_texts = [c.text.strip() for c in table.rows[0].cells] if table.rows else []
        if ncols == 12:
            _apply_col_widths(table, _COL_WIDTHS_12)
        elif ncols == 5 and "Bảng tham chiếu" in header_texts:
            # FK table (5 cột: Tên trường | Tên cột | Bảng TK | Trường TK | Cột TK)
            _apply_col_widths(table, _COL_WIDTHS_FK)
        elif ncols == 4:
            if header_texts and header_texts[0] == "STT" and "Thực thể" in header_texts:
                _apply_col_widths(table, _COL_WIDTHS_4)
        elif ncols == 2:
            # PK table
            _apply_col_widths(table, _COL_WIDTHS_PK)

        for row_idx, row in enumerate(table.rows):
            is_header = row_idx == 0
            for cell in row.cells:
                # Xóa cell-level margin override
                tcPr = cell._tc.find(qn("w:tcPr"))
                if tcPr is not None:
                    for old in tcPr.findall(qn("w:tcMar")):
                        tcPr.remove(old)

                for para in cell.paragraphs:
                    pPr = para._p.get_or_add_pPr()
                    # Xóa ind, tabs, numPr thừa
                    for ind in pPr.findall(qn("w:ind")):
                        pPr.remove(ind)
                    for tabs in pPr.findall(qn("w:tabs")):
                        pPr.remove(tabs)
                    for numPr in pPr.findall(qn("w:numPr")):
                        pPr.remove(numPr)
                    # Explicit zero indent — override docDefaults và style inheritance (kể cả w:hanging)
                    ind_el = OxmlElement("w:ind")
                    ind_el.set(qn("w:left"), "0")
                    ind_el.set(qn("w:firstLine"), "0")
                    pPr.append(ind_el)
                    # spacing before/after = 0
                    spacing = pPr.find(qn("w:spacing"))
                    if spacing is None:
                        spacing = OxmlElement("w:spacing")
                        pPr.append(spacing)
                    spacing.set(qn("w:before"), "0")
                    spacing.set(qn("w:after"), "0")
                    # Override paragraph style về Normal để tránh Body Text indent
                    pStyle = pPr.find(qn("w:pStyle"))
                    if pStyle is not None:
                        pStyle.set(qn("w:val"), "Normal")
                    # font size qua rPr default của paragraph
                    rPrDef = pPr.find(qn("w:rPr"))
                    if rPrDef is None:
                        rPrDef = OxmlElement("w:rPr")
                        pPr.append(rPrDef)
                    for tag in ("w:sz", "w:szCs"):
                        el = rPrDef.find(qn(tag))
                        if el is None:
                            el = OxmlElement(tag)
                            rPrDef.append(el)
                        el.set(qn("w:val"), "18")  # 18 half-points = 9pt
                    if is_header:
                        for tag in ("w:b", "w:bCs"):
                            el = rPrDef.find(qn(tag))
                            if el is None:
                                el = OxmlElement(tag)
                                rPrDef.insert(0, el)
                            el.attrib.pop(qn("w:val"), None)
                        # Tắt italic kế thừa từ Heading style
                        for tag in ("w:i", "w:iCs"):
                            el = rPrDef.find(qn(tag))
                            if el is None:
                                el = OxmlElement(tag)
                                rPrDef.append(el)
                            el.set(qn("w:val"), "0")
                    # font size + bold/italic cho từng run
                    for run in para.runs:
                        run.font.size = Pt(9)
                        if is_header:
                            run.font.bold = True
                            rPr = run._r.get_or_add_rPr()
                            for tag in ("w:b", "w:bCs"):
                                el = rPr.find(qn(tag))
                                if el is None:
                                    el = OxmlElement(tag)
                                    rPr.insert(0, el)
                                el.attrib.pop(qn("w:val"), None)
                            for tag in ("w:i", "w:iCs"):
                                el = rPr.find(qn(tag))
                                if el is None:
                                    el = OxmlElement(tag)
                                    rPr.append(el)
                                el.set(qn("w:val"), "0")

    doc.save(str(docx_path))
    print(f"Post-process done (landscape A4, col widths, font 9pt, header bold): {docx_path}", file=sys.stderr)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("--mode", required=True, choices=["sample", "full", "docx"])
    parser.add_argument("--source", help="Source system code (FIMS, IDS, ...). Bắt buộc cho mode sample/full.")
    parser.add_argument("--sample-count", type=int, default=2, help="Số entity trong sample mode (default 2)")
    args = parser.parse_args()

    if args.mode in ("sample", "full") and not args.source:
        parser.error(f"--source bắt buộc cho mode {args.mode}")

    if args.mode == "sample":
        out = cmd_sample(args.source, args.sample_count)
        print(f"OK sample: {out}")
    elif args.mode == "full":
        out = cmd_full(args.source)
        print(f"OK fragment: {out}")
    elif args.mode == "docx":
        out = cmd_docx(source_filter=args.source)
        print(f"OK DOCX: {out}")


if __name__ == "__main__":
    main()
